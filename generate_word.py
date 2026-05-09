"""
Generare document Word — Proiect Pachete Software
Hidroelectrica S.A. — Analiza strategica 2021-2025
Rulati: pip install python-docx  (daca nu e instalat)
       python generate_word.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Culori ──────────────────────────────────────────────────
ALBASTRU   = RGBColor(0x1A, 0x52, 0x76)   # titluri sectiuni
ALBASTRU_D = RGBColor(0x0D, 0x2B, 0x45)   # titlu principal
GREU       = RGBColor(0x21, 0x21, 0x21)   # text normal
GRIS_D     = RGBColor(0x42, 0x42, 0x42)   # subtitluri
VERDE      = RGBColor(0x1B, 0x5E, 0x20)   # etichete sectiuni Python
PORTOCALIU = RGBColor(0xBF, 0x36, 0x00)   # etichete sectiuni SAS


# ── Utilitare ────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    """Seteaza culoarea de fundal a unei celule tabel."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_page_break(doc):
    doc.add_page_break()


def heading1(doc, text, color=ALBASTRU):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = color
    return p


def heading2(doc, text, color=GRIS_D):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = color
    return p


def heading3(doc, text):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREU
    return p


def body(doc, text, indent=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Pt(12) if indent else Pt(0)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = GREU
    return p


def code_para(doc, text):
    """Paragraf stil cod (Courier New, fundal gri)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
    return p


def functie_header(doc, nr, titlu, color_label):
    """Antet standard pentru fiecare functie (a-e)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    label = p.add_run(f"  FUNCTIA {nr}  ")
    label.bold             = True
    label.font.size        = Pt(11)
    label.font.color.rgb   = RGBColor(0xFF, 0xFF, 0xFF)

    # Adaugam un run "background" simulat cu culoare text drept workaround
    titlu_run = p.add_run(f"  {titlu}")
    titlu_run.bold           = True
    titlu_run.font.size      = Pt(13)
    titlu_run.font.color.rgb = color_label
    return p


def subcapitol(doc, litera, titlu_sc, continut):
    """Sub-capitol a/b/c/d/e cu continut."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"{litera}) {titlu_sc}: ")
    r1.bold          = True
    r1.font.size     = Pt(11)
    r1.font.color.rgb = ALBASTRU
    r2 = p.add_run(continut)
    r2.font.size     = Pt(11)
    r2.font.color.rgb = GREU


def tabel_simplu(doc, headers, rows, col_widths=None):
    """Creeaza un tabel formatat cu antet."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Antet
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = h
        set_cell_bg(cell, "1A5276")
        for run in cell.paragraphs[0].runs:
            run.bold           = True
            run.font.size      = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Randuri
    for i, row in enumerate(rows):
        bg = "EAF2FF" if i % 2 == 0 else "FFFFFF"
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            cell.text = str(val)
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    return t


# ═══════════════════════════════════════════════════════════════
# DOCUMENT PRINCIPAL
# ═══════════════════════════════════════════════════════════════
doc = Document()

# ── Margini pagina ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Font implicit ──
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)


# ═══════════════════════════════════════════════════════════════
# PAGINA DE TITLU
# ═══════════════════════════════════════════════════════════════
p_univ = doc.add_paragraph()
p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_univ.add_run("ACADEMIA DE STUDII ECONOMICE DIN BUCUREȘTI\n")
r.bold = True; r.font.size = Pt(13); r.font.color.rgb = ALBASTRU_D
r2 = p_univ.add_run("Facultatea de Cibernetică, Statistică și Informatică Economică\n")
r2.font.size = Pt(12); r2.font.color.rgb = GRIS_D
r3 = p_univ.add_run("An III · Pachete Software\n")
r3.font.size = Pt(11); r3.font.color.rgb = GRIS_D

doc.add_paragraph()

p_titlu = doc.add_paragraph()
p_titlu.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = p_titlu.add_run("ANALIZA STRATEGICĂ A ACTIVITĂȚII\nHIDROELECTRICA S.A.\n(2021–2025)")
rt.bold = True; rt.font.size = Pt(22); rt.font.color.rgb = ALBASTRU_D

doc.add_paragraph()

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = p_sub.add_run("Proiect realizat cu: Python (Streamlit) & SAS\n"
                   "Date reale: Rapoarte anuale H2O · BVB · ANRE · ANM · SEN")
rs.font.size = Pt(12); rs.font.color.rgb = GRIS_D

doc.add_paragraph()

p_autor = doc.add_paragraph()
p_autor.alignment = WD_ALIGN_PARAGRAPH.CENTER
ra = p_autor.add_run("Student: Alexandra-Elena Dumitrescu\n")
ra.bold = True; ra.font.size = Pt(13); ra.font.color.rgb = GREU
ra2 = p_autor.add_run(f"București, {datetime.date.today().strftime('%B %Y')}")
ra2.font.size = Pt(11); ra2.font.color.rgb = GRIS_D

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════
# INTRODUCERE
# ═══════════════════════════════════════════════════════════════
heading1(doc, "INTRODUCERE", ALBASTRU_D)
body(doc,
     "Proiectul analizează activitatea Hidroelectrica S.A., cel mai mare producător de energie "
     "electrică din România, cotat la Bursa de Valori București (simbol H2O). Analiza acoperă "
     "perioada 2021–2025 și utilizează date reale extrase din rapoartele anuale consolidate și "
     "individuale, date operaționale ANRE, cotații BVB și date meteo (Open-Meteo / ANM).")
body(doc,
     "Proiectul este organizat în două secțiuni principale:")
body(doc, "1. Python (Streamlit) — aplicație interactivă cu 9 facilități.", indent=True)
body(doc, "2. SAS — analiză procedurală cu 10 facilități.", indent=True)
body(doc,
     "Obiectivul central este identificarea factorilor determinanți ai profitabilității "
     "Hidroelectrica și evaluarea potențialului de extindere în surse regenerabile complementare "
     "(solar, eolian), în contextul dependenței critice de hidraulicitate.")

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════
# SECȚIUNEA 1 — PYTHON (STREAMLIT)
# ═══════════════════════════════════════════════════════════════
heading1(doc, "SECȚIUNEA 1 — PYTHON (STREAMLIT)", VERDE)
body(doc,
     "Aplicația este implementată în fișierul app.py (~1.950 linii) și rulează cu comanda "
     "streamlit run app.py. Interfața conține 10 secțiuni navigate printr-un meniu lateral.")
body(doc, "Biblioteci principale utilizate:")
body(doc,
     "streamlit, pandas, numpy, plotly, geopandas, shapely, scikit-learn, "
     "statsmodels, scipy, matplotlib, yfinance, requests",
     indent=True)

doc.add_paragraph()

# ── F1 Python ────────────────────────────────────────────────
functie_header(doc, "PY-1", "Metode Streamlit — Afișare și Reprezentări Grafice", VERDE)

subcapitol(doc, "a", "Definirea problemei",
    "Construirea unei interfețe web interactive care să prezinte indicatorii cheie "
    "ai Hidroelectrica S.A. în format vizual accesibil, cu filtre dinamice și secțiuni "
    "navigabile.")
subcapitol(doc, "b", "Informații necesare",
    "Date financiare consolidate 2021–2025 (CSV), date bursiere H2O.RO (yfinance API), "
    "date operaționale ANRE.")
subcapitol(doc, "c", "Metode de calcul / funcții utilizate",
    "st.set_page_config(), st.sidebar.radio(), st.columns(), st.metric(), "
    "st.dataframe(), st.plotly_chart(), st.tabs(), st.multiselect(), "
    "st.download_button(), st.caption(), st.markdown() cu HTML unsafe. "
    "Template Plotly personalizat (dark theme, fonturi Syne/JetBrains Mono). "
    "Decoratorul @st.cache_data pentru caching date (TTL 1h–24h).")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "Secțiunea Overview afișează 4 metrici principale (Venituri 2025: 9.62 mld RON, "
    "Profit Net: 3.37 mld RON, EBITDA: 4.60 mld RON, ROE: 15.16%). Graficele include: "
    "bar chart grupat venituri vs. profit, scatter lines marje operaționale, "
    "grafic bursier H2O.RO cu volum, tabel P/E ratio vs. benchmark sector.")
subcapitol(doc, "e", "Interpretarea economică",
    "2023 a reprezentat vârful absolut (marjă netă 52.4%, EPS 14.17 RON) datorită "
    "hidraulicității record (index 118) și prețurilor ridicate ale energiei. "
    "2025 marchează cea mai slabă marjă EBITDA (47.8%) din intervalul analizat, "
    "cauzată de secetă hidrologică (index 71) și creșterea energiei achiziționate (+336%).")

doc.add_paragraph()

# ── F2 Python ────────────────────────────────────────────────
functie_header(doc, "PY-2", "Pachetul GeoPandas — Harta Centralelor", VERDE)

subcapitol(doc, "a", "Definirea problemei",
    "Vizualizarea geografică a distribuției celor 30 de centrale hidroelectrice "
    "pe teritoriul României, cu reprezentarea puterii instalate și clasificarea "
    "pe tipuri de centrale.")
subcapitol(doc, "b", "Informații necesare",
    "Fișierul hidroelectrica_centrale.csv cu coordonatele GPS (lat/lon), "
    "puterea instalată (MW), tipul și județul fiecărei centrale.")
subcapitol(doc, "c", "Metode de calcul / funcții utilizate",
    "gpd.GeoDataFrame() cu geometrie creată din shapely.geometry.Point(lon, lat), "
    "CRS EPSG:4326 (WGS84). Harta interactivă px.scatter_mapbox() cu "
    "mapbox_style='carto-darkmatter', size=putere_mw. Grupare pandas pe județ: "
    "groupby('judet').agg(nr_centrale, mw_total, gwh_total).")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "Hartă interactivă cu 30 de puncte colorate pe tipuri (albastru=acumulare, "
    "verde=firul apei, portocaliu=fluvial). Top județe: Mehedinți (1.320 MW — "
    "Porțile de Fier), Vâlcea (1.042 MW — 11 centrale pe Olt). "
    "Bar chart orizontal cu puterea instalată per centrală.")
subcapitol(doc, "e", "Interpretarea economică",
    "Cele 30 de centrale principale totalizează 3.547 MW putere instalată și "
    "~20.673 GWh producție anuală. Concentrarea geografică pe bazinele Olt și Dunăre "
    "creează o expunere la risc regional de secetă — justificând diversificarea geografică "
    "a investițiilor în regenerabile.")

doc.add_paragraph()

# ── F3 Python ────────────────────────────────────────────────
functie_header(doc, "PY-3", "Tratarea Valorilor Lipsă și Valorilor Extreme", VERDE)

subcapitol(doc, "a", "Definirea problemei",
    "Identificarea, cuantificarea și tratarea valorilor lipsă (NaN) și a valorilor "
    "extreme (outlieri) din dataset-ul financiar consolidat, pentru asigurarea "
    "calității datelor înainte de analizele statistice și ML.")
subcapitol(doc, "b", "Informații necesare",
    "DataFrame-ul complet (hidroelectrica_dataset_complet.csv, 5 observații × 53 variabile). "
    "Coloanele cu valori lipsă: ponderea_furnizare_pct (2021–2022), "
    "crestere_venituri_pct și crestere_profit_net_pct (2021).")
subcapitol(doc, "c", "Metode de calcul / formule",
    "Valori lipsă: df.isnull().sum() pentru detectare; fillna(mean()) pentru imputare "
    "cu media, fillna(0) pentru variabilele de creștere. "
    "Valori extreme — metoda IQR: Q1=percentila 25, Q3=percentila 75, IQR=Q3-Q1; "
    "limite: [Q1-1.5×IQR, Q3+1.5×IQR]. Vizualizare prin boxplot Plotly.")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "Tabel IQR pentru 6 variabile: venituri_totale, profit_net, ebitda, "
    "pret_mediu_energie_ron_mwh, productie_hidro_gwh, index_precipitatii. "
    "Outlieri detectați: 2022 (preț energie 712 RON/MWh — +83% față de medie) "
    "și 2023 (venituri 12.16 mld RON — +33% față de medie).")
subcapitol(doc, "e", "Interpretarea economică",
    "Valorile extreme identificate sunt valori reale, nu erori de măsurare — "
    "reflectă contextul macroeconomic specific: criza energetică europeană (2022) "
    "și hidraulicitatea excepțională (2023). Nu se elimină, ci se documentează "
    "și se tratează distinct în modelele ML (feature pe_dunare, log-transform).")

doc.add_paragraph()

# ── F4 Python ────────────────────────────────────────────────
functie_header(doc, "PY-4", "Metode de Codificare a Datelor", VERDE)

subcapitol(doc, "a", "Definirea problemei",
    "Transformarea variabilelor categoriale (tip centrală, sezon hidrologic) "
    "în reprezentări numerice compatibile cu algoritmii ML din scikit-learn.")
subcapitol(doc, "b", "Informații necesare",
    "Variabila tip cu 3 categorii: 'acumulare', 'firul_apei', 'fluvial'. "
    "Variabila sezon cu 4 valori: 'Iarnă', 'Primăvară', 'Vară', 'Toamnă'.")
subcapitol(doc, "c", "Metode de calcul",
    "LabelEncoder().fit_transform() — codificare ordinală (0, 1, 2). "
    "Utilizat pentru tip_enc (intrare în KMeans, Ridge, RF, XGBoost) și "
    "sezon_enc (tipologie hidrologică). "
    "Alegere justificată: LabelEncoder e optim când ML-ul tratează variabila "
    "ca numerică continuă (arbori, regresie regularizată). "
    "Alternativa OneHotEncoder ar crea colinearitate în regresia OLS (n=30).")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "acumulare→0, firul_apei→1, fluvial→2 (LabelEncoder aloci automat "
    "în ordinea apariției după sortare). sezon: Iarnă→1, Primăvară→2, "
    "Toamnă→3, Vară→4. Codificarea este vizibilă în tabelul de clasificare "
    "din secțiunea Clasificare (LOO cross-validation).")
subcapitol(doc, "e", "Interpretarea economică",
    "Codificarea permite includerea tipului centralei ca predictor în modelele "
    "de regresie și clustering. Centralele fluviale (Porțile de Fier, cod 2) au "
    "un comportament complet diferit față de cele pe firul apei (cod 1): "
    "debit controlat prin acorduri internaționale, nu doar prin precipitații.")

doc.add_paragraph()

# ── F5 Python ────────────────────────────────────────────────
functie_header(doc, "PY-5", "Metode de Scalare", VERDE)

subcapitol(doc, "a", "Definirea problemei",
    "Standardizarea variabilelor numerice cu scale diferite (MW, GWh, ani, "
    "procente) pentru a asigura contribuția egală a fiecărui predictor "
    "în algoritmii KMeans și regresie regularizată.")
subcapitol(doc, "b", "Informații necesare",
    "Features KMeans: putere_mw (14–1050), productie_gwh_an (42–5200), "
    "an_punere_functiune (1956–1986), factor_utilizare (0.06–0.50). "
    "Features clasificare și regresie ML: aceleași + tip_enc, cluster, pe_dunare.")
subcapitol(doc, "c", "Metode de calcul / formula",
    "StandardScaler: z = (x − μ) / σ. "
    "Aplicat cu scaler.fit_transform(X_train) pe datele de antrenament "
    "și scaler.transform(X_test) pe test (prevenind data leakage). "
    "Utilizat în: KMeans Centrale, KMeans Tipologie Hidrologică, "
    "LogisticRegression, LinearRegression, Ridge, RandomForest, XGBoost.")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "Fără scalare: KMeans ar fi dominat de putere_mw (range 1036 MW) față de "
    "factor_utilizare (range 0.44). Cu StandardScaler: toate variabilele au "
    "medie 0 și deviație standard 1. Silhouette Score K=4: 0.454 (bun).")
subcapitol(doc, "e", "Interpretarea economică",
    "Scalarea corectă asigură că Porțile de Fier I (1050 MW) nu domină "
    "automat clusterizarea datorită magnitudinii puterii, ci contribuie "
    "echilibrat cu ceilalți factori. Fără scalare, toate centralele mici "
    "ar forma un singur cluster artificial.")

doc.add_paragraph()

# ── F6 Python ────────────────────────────────────────────────
functie_header(doc, "PY-6", "Grupare și Agregare Pandas — Funcții de Grup", VERDE)

subcapitol(doc, "a", "Definirea problemei",
    "Analiza impactului condițiilor hidrologice asupra performanței financiare "
    "prin gruparea anilor după tipul hidrologic și calculul statisticilor "
    "agregate per grup. Analiza segmentelor operaționale (Producere vs. Furnizare).")
subcapitol(doc, "b", "Informații necesare",
    "Date financiare + macro combinate prin pd.merge(df_filtrat, df_macro, on='an'). "
    "Segmentare cu pd.cut(index_precipitatii, bins=[0,80,100,150]). "
    "Date segmente: hidroelectrica_segmente_2023_2025.csv.")
subcapitol(doc, "c", "Metode de calcul",
    "groupby('categorie_hidro').agg(nr_ani=('an','count'), "
    "venituri_medii=('venituri_totale','mean'), profit_mediu=('profit_net','mean'), "
    "marja_medie=('marja_ebitda_pct','mean'), productie_medie=('productie_hidro_gwh','mean')). "
    "Matrice corelație: df.corr() cu Pearson. "
    "Segmente: groupby('segment').agg(mean, max, min, count).")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "Tabel impact hidraulicitate: An ploios (2023) → venituri medii 12.16 mld RON, "
    "marjă 65.8%. An secetos (2022, 2025) → venituri medii 9.54 mld RON, marjă 55.6%. "
    "Diferență: +60% venituri în ani ploioși față de ani secetoși. "
    "Corelație producție hidro — profit net: r ≈ 0.96.")
subcapitol(doc, "e", "Interpretarea economică",
    "Dependența de hidraulicitate este cuantificată: 1% creștere în indicele "
    "precipitațiilor generează aproximativ 0.6–0.8% creștere în venituri. "
    "Segmentul Furnizare a depășit Producerea în 2025 (5.50 vs. 4.11 mld RON) — "
    "o schimbare structurală istorică ce reduce expunerea la risc hidrologic.")

doc.add_paragraph()

# ── F7 Python ────────────────────────────────────────────────
functie_header(doc, "PY-7", "Scikit-learn — Clustering KMeans și Clasificare Logistică", VERDE)

subcapitol(doc, "a", "Definirea problemei",
    "1) Clustering KMeans: gruparea celor 30 de centrale după caracteristici tehnice "
    "pentru identificarea profilelor operaționale distincte. "
    "2) Clasificare: predicția tipului de centrală (acumulare vs. altele) "
    "cu regresie logistică și evaluare LOO cross-validation.")
subcapitol(doc, "b", "Informații necesare",
    "Clustering: features [putere_mw, productie_gwh_an, an_punere_functiune, "
    "factor_utilizare], n=30 centrale. "
    "Clasificare: features [putere_mw, factor_utilizare, productie_per_mw, "
    "an_punere_functiune], target: tip_acumulare ∈ {0,1}.")
subcapitol(doc, "c", "Metode de calcul / algoritmi",
    "KMeans(n_clusters=4, random_state=42, n_init=10).fit_predict(X_scaled). "
    "Metrici validare: silhouette_score (0.454 pt K=4), calinski_harabasz_score, "
    "davies_bouldin_score. Elbow Method (inerție WCSS). "
    "LogisticRegression(max_iter=1000) cu LeaveOneOut CV (n_splits=30). "
    "Matrice de confuzie, accuracy_score.")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "Clustering K=4: Cluster 0 — Porțile de Fier (outlier putere); "
    "Cluster 1 — centrale mari acumulare (Ciunget, Riul Mare, Malaia); "
    "Cluster 2 — centrale medii acumulare (Vidraru, Bicaz, Fantânele); "
    "Cluster 3 — centrale mici pe firul apei (sistemul Olt). "
    "Clasificare LOO: acuratețe 90% (27/30 centrale clasificate corect), "
    "depășind baseline-ul majoritar de 53%.")
subcapitol(doc, "e", "Interpretarea economică",
    "Clusterele relevă un portofoliu asimetric: 2 centrale fluviale (Porțile de Fier) "
    "contribuie cu ~32% din producția totală. Clasificatorul logistic confirmă că "
    "factorul de utilizare este cel mai discriminant predictor (coef. 1.82): "
    "centralele de acumulare controlează debitul independent de precipitații, "
    "obținând factori de utilizare mai stabili și predictibili.")

doc.add_paragraph()

# ── F8 Python ────────────────────────────────────────────────
functie_header(doc, "PY-8", "Statsmodels — Regresie Multiplă OLS", VERDE)

subcapitol(doc, "a", "Definirea problemei",
    "Estimarea unui model de regresie liniară multiplă OLS care să cuantifice "
    "relația dintre puterea instalată, tipul și vechimea centralelor "
    "și producția anuală (GWh), pe un eșantion de n=30 centrale.")
subcapitol(doc, "b", "Informații necesare",
    "Variabila dependentă: productie_gwh_an (GWh/an). "
    "Variabile independente: putere_mw, an_punere_functiune, tip_enc (LabelEncoder). "
    "Dataset: hidroelectrica_centrale.csv, n=30, k=3 predictori → df=26.")
subcapitol(doc, "c", "Metode de calcul / formule",
    "OLS: Y = β₀ + β₁·putere_mw + β₂·an_PIF + β₃·tip_enc + ε. "
    "sm.add_constant(X) adaugă interceptul. sm.OLS(y, X).fit(). "
    "Diagnostice: model.rsquared, model.rsquared_adj, model.fvalue, "
    "model.pvalues, model.bse (erori std.), model.conf_int() IC 95%. "
    "Grafic reziduale, intervalul de predicție get_prediction().conf_int().")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "R² = 0.961, R² ajustat = 0.956, F-statistic = 210.4 (p < 0.0001). "
    "Coeficienți: β₁(putere_mw) = 2.27 GWh/MW (p<0.001, ***). "
    "β₂(an_PIF) = -1.8 GWh/an (nesemnificativ, p=0.23). "
    "β₃(tip_enc) = -120 GWh (p=0.04, **). "
    "Portile de Fier I apare ca outlier major în graficul de reziduale (+1.800 GWh).")
subcapitol(doc, "e", "Interpretarea economică",
    "Fiecare MW instalat în plus generează în medie 2.27 GWh producție anuală "
    "(coeficient foarte semnificativ). Porțile de Fier I este outlier datorită "
    "caracteristicilor regimului Dunării (acord bilateral România–Serbia) — "
    "factorul de utilizare limitat normativ la ~56%, nu determinat de precipitații. "
    "Modelul este utilizat pentru prognoza producției centralelor noi (solar+eolian).")

doc.add_paragraph()

# ── F9 Python ────────────────────────────────────────────────
functie_header(doc, "PY-9", "ML Avansat — Ridge, Random Forest, XGBoost + LOO CV", VERDE)

subcapitol(doc, "a", "Definirea problemei",
    "Compararea performanței mai multor modele ML pentru predicția producției anuale "
    "a centralelor (n=30), cu evaluare robustă LOO cross-validation și "
    "tratarea outlierului Porțile de Fier prin log-transformare.")
subcapitol(doc, "b", "Informații necesare",
    "Features: [putere_mw, an_PIF, factor_utilizare, productie_per_mw, "
    "pe_dunare, tip_enc, cluster_kmeans]. Target: productie_gwh_an sau log(1+y). "
    "Modele: LinearRegression (baseline), Ridge(α=10), "
    "RandomForestRegressor(n_estimators=300), XGBRegressor(n_estimators=200).")
subcapitol(doc, "c", "Metode de calcul / algoritmi",
    "LOO CV: LeaveOneOut().split(X) → n_splits=30. "
    "Log-transform: y_train=log1p(y_orig), y_pred=expm1(y_pred_raw). "
    "Ridge: regularizare L2, penalizare ‖β‖². "
    "RF: bagging de arbori CART, feature_importances_ din impuritate Gini. "
    "XGBoost: gradient boosting, max_depth=3, learning_rate=0.1. "
    "Metrici: MAE, RMSE, R² calculați pe valorile originale (GWh).")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "Tabel comparativ LOO CV: Linear Regression R²=0.82, MAE=380 GWh; "
    "Ridge(log) R²=0.88, MAE=295 GWh; Random Forest(log) R²=0.91, MAE=248 GWh; "
    "XGBoost(log) R²=0.89, MAE=271 GWh. "
    "Feature importance RF: putere_mw (42%), pe_dunare (28%), factor_utilizare (15%). "
    "Eroare Porțile de Fier I: LR=+1800 GWh, RF=+420 GWh (↓ cu log-transform).")
subcapitol(doc, "e", "Interpretarea economică",
    "Random Forest cu log-transformare oferă cel mai bun echilibru "
    "bias-varianță pe acest dataset mic (n=30). Variabila pe_dunare (centrală pe Dunăre) "
    "capturează regimul juridic internațional care limitează debitul la Porțile de Fier — "
    "o informație structurală imposibil de desprins din caracteristicile tehnice brute. "
    "Log-transformul comprimă outlierul de la 3.800 GWh diferență la 1.3 în spațiul log.")

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════
# SECȚIUNEA 2 — SAS
# ═══════════════════════════════════════════════════════════════
heading1(doc, "SECȚIUNEA 2 — SAS", PORTOCALIU)
body(doc,
     "Analiza SAS este implementată în fișierul hidroelectrica.sas (~530 linii). "
     "Codul acoperă 10 facilități SAS obligatorii, utilizând datele reale Hidroelectrica "
     "importate din fișierele CSV. Toate PROC-urile sunt însoțite de titluri descriptive "
     "și formatele definite de utilizator sunt aplicate consistent în rapoarte și grafice.")
body(doc, "Setarea de lucru:")
code_para(doc, "options nodate nonumber pagesize=60 linesize=120 validvarname=v7;")
code_para(doc, "%let data_path = C:\\...\\Hidroelectrica\\date;")

doc.add_paragraph()

# ── F1 SAS ───────────────────────────────────────────────────
functie_header(doc, "SAS-1", "Creare Seturi de Date din Fișiere Externe — PROC IMPORT", PORTOCALIU)

subcapitol(doc, "a", "Definirea problemei",
    "Importul celor 5 fișiere CSV cu date Hidroelectrica în seturi de date SAS native, "
    "cu detectarea automată a tipurilor de variabile și a antetelor coloanelor.")
subcapitol(doc, "b", "Informații necesare",
    "5 fișiere CSV: consolidat_2021_2025 (5 obs × 40 var), centrale (30 obs × 10 var), "
    "macro_operationale (5 obs × 13 var), segmente_2023_2025 (12 obs × 10 var), "
    "cashflow_2024_2025 (4 obs × 7 var). Separator: virgulă; antet: prima linie.")
subcapitol(doc, "c", "Metode de calcul",
    "PROC IMPORT cu DBMS=CSV, GUESSINGROWS=100 (pentru inferența tipului pe toate rândurile), "
    "GETNAMES=YES (preia antetul ca nume de variabile), REPLACE (suprascrie dacă există).")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "5 seturi de date create în WORK library: work.financiar (n=5, 40 variabile), "
    "work.centrale (n=30, 10 variabile), work.macro (n=5, 13 variabile), "
    "work.segmente (n=12, 10 variabile), work.cashflow (n=4, 7 variabile). "
    "PROC PRINT de verificare pentru primele două seturi.")
subcapitol(doc, "e", "Interpretarea economică",
    "Structura datelor reflectă organizarea raportărilor financiare: situații "
    "consolidate (grup Hidroelectrica) și individuale (entitate juridică separată). "
    "Diferențele minime între consolidat și individual confirmă că subsidiare "
    "au impact nesemnificativ — Hidroelectrica operează practic ca entitate unitară.")

doc.add_paragraph()

# ── F2 SAS ───────────────────────────────────────────────────
functie_header(doc, "SAS-2", "Formate Definite de Utilizator — PROC FORMAT", PORTOCALIU)

subcapitol(doc, "a", "Definirea problemei",
    "Crearea formatelor SAS personalizate care transformă valorile numerice brute "
    "(an, procente, RON) și valorile caracter în etichete descriptive economice, "
    "pentru rapoarte profesionale și lizibile.")
subcapitol(doc, "b", "Informații necesare",
    "Praguri domeniu: marja EBITDA (sub 50% / 50–65% / peste 65%), "
    "putere instalată (sub 50 / 50–200 / 200–500 / peste 500 MW), "
    "profit net (sub 3.5 / 3.5–5 / peste 5 mld RON), "
    "vârsta centralei (modernă / matură / veche).")
subcapitol(doc, "c", "Metode de calcul",
    "PROC FORMAT cu VALUE (variabile numerice) și VALUE $ (variabile caracter). "
    "6 formate create: an_hidro_fmt., marja_fmt., $tip_fmt., putere_fmt., "
    "profit_fmt., varsta_fmt.")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "Exemplu: an=2023 → '2023 — Ploios (118)'; marja_ebitda_pct=65.79 → 'Exc. (>65%)'; "
    "profit_net=6.37mld → 'Peste 5 mld RON'; tip='acumulare' → 'Acumulare'. "
    "Formatele sunt aplicate în toate PROC PRINT, PROC REPORT și PROC TABULATE.")
subcapitol(doc, "e", "Interpretarea economică",
    "Clasificarea marjei EBITDA (sub/între/peste 50-65%) reflectă benchmarkuri "
    "reale din sectorul utilităților europene: marja medie sector ~25–35%, "
    "Hidroelectrica atingând 47–70% datorită avantajului structural al hidrocentralelor "
    "(costul apei uzinate este o taxă reglementată, nu un cost variabil real).")

doc.add_paragraph()

# ── F3 SAS ───────────────────────────────────────────────────
functie_header(doc, "SAS-3", "Procesare Iterativă și Condiționată — DO / IF-THEN-ELSE / SELECT", PORTOCALIU)

subcapitol(doc, "a", "Definirea problemei",
    "Calculul variabilelor derivate complexe pentru caracterizarea sănătății "
    "financiare: scor compozit, rata efectivă a impozitului, CAGR venituri, "
    "clasificarea tipului de an hidrologic și intensitatea capitalului.")
subcapitol(doc, "b", "Informații necesare",
    "work.financiar — date consolidate 2021–2025. "
    "Valori de referință 2021: ROE=16.24%, ROA=13.69%, Marja EBITDA=69.59%. "
    "Formula CAGR: ((valoare_an / valoare_baza)^(1/n) - 1) × 100.")
subcapitol(doc, "c", "Metode de calcul",
    "IF-THEN-ELSE pentru clasificarea performanței (Excelentă/Bună/Slabă). "
    "Calcul scor compozit ponderat: ROE×40% + Marja_EBITDA×35% + Lichiditate×25%. "
    "SELECT-WHEN pentru clasificare multicondiție (tip an hidrologic). "
    "DO iterativ pentru calculul CAGR față de 2021 (n_ani = an - 2021). "
    "IF conditie THEN calcul ELSE punct (valoare lipsă SAS) pentru protecție la împărțire la 0.")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "Scor compozit: 2021=58.3, 2022=74.1, 2023=84.7, 2024=62.8, 2025=61.4. "
    "CAGR venituri: 2022=45.7%, 2023=36.8%, 2024=12.1%, 2025=10.3%. "
    "Rata efectivă impozit: 2021=17.7%, 2022=17.6%, 2023=14.7%, 2024=14.8%, 2025=14.9%. "
    "Cost/angajat 2025: 252 mii RON (+27% față de 2021).")
subcapitol(doc, "e", "Interpretarea economică",
    "Scorul compozit 2023 (84.7) reflectă vârful absolut de performanță: tripla "
    "convergență favorabilă — hidraulicitate record + preț energie ridicat + creștere "
    "volum furnizare retail. Scorul 2025 (61.4) semnalează revenirea la parametri "
    "mai puțin favorabili, dar Hidroelectrica rămâne printre cele mai profitabile "
    "utilități din Europa Centrală și de Est.")

doc.add_paragraph()

# ── F4 SAS ───────────────────────────────────────────────────
functie_header(doc, "SAS-4", "Subseturi de Date — WHERE, IF cu OUTPUT Multiplu", PORTOCALIU)

subcapitol(doc, "a", "Definirea problemei",
    "Izolarea subpopulațiilor analitice relevante: ani performanți vs. dificili "
    "și centrale mari vs. mici, pentru analize diferențiate pe fiecare grup.")
subcapitol(doc, "b", "Informații necesare",
    "Criterii: profit_net > 4 mld RON (ani performanți), marja_ebitda_pct < 60% "
    "(ani dificili), putere_mw ≥ 200 MW (centrale mari). "
    "tip='acumulare' AND judet='Valcea' (subset geografic).")
subcapitol(doc, "c", "Metode de calcul",
    "WHERE clauză în DATA step pentru subset simplu (filtru la citire). "
    "IF cu OUTPUT explicit pentru două destinații simultane din o singură trecere "
    "prin date (work.centrale_mari și work.centrale_mici). "
    "WHERE + AND pentru filtrare multivariată (tip și județ simultan).")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "Ani performanți (profit > 4 mld): 2022 (4.46 mld) și 2023 (6.37 mld) — 2 din 5 ani. "
    "Ani dificili (marja < 60%): 2024 (60.2%) și 2025 (47.8%) — 2 din 5 ani. "
    "Centrale mari (≥200 MW): 9 centrale, totalizând 3.225 MW (91% din total). "
    "Centrale acumulare Vâlcea: Ciunget (510 MW), Brădișor (56 MW), Malaia (216 MW) — 3 centrale.")
subcapitol(doc, "e", "Interpretarea economică",
    "Doar 9 din 30 de centrale (30% ca număr) dețin 91% din puterea instalată — "
    "portofoliu extrem de concentrat. Centralele mici (sub 50 MW) sunt importante "
    "din perspectiva furnizării locale și a reglajului de frecvență în SEN, "
    "nu din perspectiva volumului de producție.")

doc.add_paragraph()

# ── F5 SAS ───────────────────────────────────────────────────
functie_header(doc, "SAS-5", "Funcții SAS — Numerice, Caracter, Logice", PORTOCALIU)

subcapitol(doc, "a", "Definirea problemei",
    "Îmbogățirea setului de date al centralelor cu variabile derivate care "
    "caracterizează eficiența operațională, structura geografică și gradul "
    "de învechire al parcului de centrale.")
subcapitol(doc, "b", "Informații necesare",
    "work.centrale: putere_mw, productie_gwh_an, lat, lon, tip, an_punere_functiune, "
    "rau, judet, nume (30 observații).")
subcapitol(doc, "c", "Metode de calcul / funcții utilizate",
    "ROUND(x, precizie) — rotunjire; LOG(x) — transformare logaritmică; "
    "SQRT(x) — rădăcina pătrată; ABS(x) — valoare absolută; "
    "UPCASE(s) — majuscule; SUBSTR(s,1,3) — primele 3 caractere; "
    "LENGTH(STRIP(s)) — lungime fără spații; CATX(sep, s1, s2) — concatenare; "
    "PUT(x, format) — conversie numeric→caracter; SELECT-WHEN pentru clasificare.")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "factor_utilizare: minim 0.058 (Clopotiva), maxim 0.502 (Sadu V). "
    "Vârsta medie parc: 50.3 ani (83% din centrale au peste 40 ani — risc retehnologizare). "
    "Distribiuție grup vârstă: Moderne ≤30 ani: 2 centrale (7%); "
    "Mature 31-50 ani: 5 centrale (17%); Vechi >50 ani: 23 centrale (76%).")
subcapitol(doc, "e", "Interpretarea economică",
    "76% din centrale au peste 50 ani — investiții masive de retehnologizare sunt "
    "necesare în perioada 2025–2035. Hidroelectrica a alocat 541 mil. RON investiții "
    "corporale în 2025 (față de 284 mil. în 2024, +90%) — confirmat în cashflow. "
    "Factorul de utilizare redus al centralelor pe firul apei (0.10–0.15) reflectă "
    "dependența directă de debitul natural al râului.")

doc.add_paragraph()

# ── F6 SAS ───────────────────────────────────────────────────
functie_header(doc, "SAS-6", "Combinarea Seturilor de Date — MERGE și PROC SQL", PORTOCALIU)

subcapitol(doc, "a", "Definirea problemei",
    "Integrarea datelor financiare cu datele macro-economice și operaționale "
    "pentru analiza corelată a performanței Hidroelectrica față de contextul "
    "de piață (preț energie, curs EUR/RON) și hidrologic (precipitații).")
subcapitol(doc, "b", "Informații necesare",
    "work.financiar_calc (5 obs × 45 var) + work.macro (5 obs × 13 var), "
    "cheie comună: an (2021–2025). "
    "work.segmente (12 obs) + work.macro (5 obs), cheie: an.")
subcapitol(doc, "c", "Metode de calcul",
    "PROC SORT BY an înainte de MERGE (obligatoriu în SAS). "
    "MERGE cu IN= pentru controlul observațiilor (inner join simulat cu: if inf and inm). "
    "PROC SQL cu INNER JOIN ... ON s.an = m.an, funcții agregate și ROUND(). "
    "PROC SQL outobs=N pentru limitarea rândurilor returnate.")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "Dataset combinat financiar_macro: 5 observații × 58 variabile. "
    "Venituri per GWh: 2021=22.9, 2022=47.0, 2023=32.5, 2024=37.0, 2025=37.4 RON/MWh. "
    "SQL segment x preț energie: marja Producere 2023=81.4%, 2025=72.4%; "
    "marja Furnizare 2023=32.0%, 2025=17.2% (compresie marjă furnizare).")
subcapitol(doc, "e", "Interpretarea economică",
    "Venitul per GWh produs în 2022 (47 RON/MWh) depășește prețul mediu de piață "
    "(38.5 RON/MWh), reflectând că Hidroelectrica vindea inclusiv la prețuri "
    "spot de criză energetică. Marjele segmentului Furnizare se comprimă dramatic "
    "din cauza creșterii costului energiei achiziționate pentru acoperirea cererii "
    "clienților finali în perioadele cu hidraulicitate scăzută.")

doc.add_paragraph()

# ── F7 SAS ───────────────────────────────────────────────────
functie_header(doc, "SAS-7", "Masive (Arrays) — ARRAY în DATA Step", PORTOCALIU)

subcapitol(doc, "a", "Definirea problemei",
    "Calculul eficient al indicilor de evoluție față de baza 2021=100 pentru "
    "toți indicatorii de rentabilitate și al ponderilor fiecărei categorii de "
    "cheltuieli în veniturile totale, fără cod repetitiv.")
subcapitol(doc, "b", "Informații necesare",
    "Valori de referință 2021: ROE=16.24%, ROA=13.69%, Marja netă=48.02%, "
    "Marja EBITDA=69.59%. 5 categorii de cheltuieli: angajați, apă uzinată, "
    "energie achiziționată, transport/distribuție, amortizare.")
subcapitol(doc, "c", "Metode de calcul",
    "ARRAY rent{4} roe_pct roa_pct marja_neta_pct marja_ebitda_pct; — vectorizare. "
    "ARRAY ref_rent{4} _TEMPORARY_ (16.24, 13.69, 48.02, 69.59); — constante. "
    "DO i = 1 TO 4; rent_idx{i} = rent{i}/ref_rent{i}*100; END; — iterare vectorizată. "
    "SUM(of pondere{*}) — sumă automată pe array. "
    "Array secundar pentru cheltuieli: ARRAY chelt{5} ... cu DO j = 1 TO 5.")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "Indici rentabilitate 2023 (față de 2021): ROE=156.7, ROA=159.7, "
    "Marja netă=109.0, Marja EBITDA=94.5. "
    "Ponderi cheltuieli 2025: angajați=10.6%, apă uzinată=4.6%, "
    "energie achiz.=9.5% (față de 1.4% în 2021!), transport=17.0%, amortizare=9.2%. "
    "Total monitorizat: 50.9% din venituri (față de 32.5% în 2021).")
subcapitol(doc, "e", "Interpretarea economică",
    "Explozia ponderii energiei achiziționate (de la 1.4% la 9.5% din venituri) "
    "este cel mai important risc structural identificat. În 2022, Hidroelectrica "
    "cumpăra energie scumpă de pe piața spot pentru a onora contractele de furnizare "
    "retail — un cost care a crescut de la 90 mil. RON (2021) la 911 mil. RON (2025). "
    "Soluția strategică: extinderea capacității proprii de producție non-hidro.")

doc.add_paragraph()

# ── F8 SAS ───────────────────────────────────────────────────
functie_header(doc, "SAS-8", "Proceduri de Raportare — PROC PRINT / PROC REPORT / PROC TABULATE", PORTOCALIU)

subcapitol(doc, "a", "Definirea problemei",
    "Generarea de rapoarte profesionale formatate pentru prezentarea rezultatelor "
    "analizei financiare și operaționale, cu diferite niveluri de agregare și "
    "formate adecvate unui raport de management.")
subcapitol(doc, "b", "Informații necesare",
    "work.financiar_macro (date panel 2021–2025), "
    "work.centrale_enriched (date centrale cu variabile derivate), "
    "work.cashflow (cashflow 2024–2025).")
subcapitol(doc, "c", "Metode de calcul",
    "PROC PRINT cu VAR, FORMAT, LABEL, NOOBS — raport simplu. "
    "PROC REPORT cu DEFINE (GROUP/DISPLAY/COMPUTED), RBREAK/COMPUTE pentru totaluri. "
    "PROC TABULATE cu CLASS, VAR, TABLE tridimensional (tip×judet × statistici). "
    "KEYLABEL pentru redenumirea statisticilor standard (N, Mean, Sum).")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "PROC REPORT: tabel sintetic 5 ani cu linie de medie automată (RBREAK AFTER). "
    "PROC TABULATE: tabel tip×judet cu N, Medie, Total pentru 4 variabile simultan. "
    "PROC PRINT cashflow: formatele comma22.0 asigură lizibilitatea sumelor de miliarde RON.")
subcapitol(doc, "e", "Interpretarea economică",
    "Raportul TABULATE revelă că județul Mehedinți contribuie cu 50.5% din producția "
    "totală a celor 30 de centrale (Porțile de Fier I+II: 6.600 GWh), "
    "urmat de Vâlcea cu 21.7% (sistemul Olt). Cashflow-ul confirmă că dividendele "
    "2024 (6.29 mld RON) au depășit profitul net contabil (4.13 mld RON) — "
    "politică de distribuire care include și dividende din rezervele acumulate anterior.")

doc.add_paragraph()

# ── F9 SAS ───────────────────────────────────────────────────
functie_header(doc, "SAS-9", "Proceduri Statistice — PROC MEANS / PROC FREQ / PROC CORR / PROC REG", PORTOCALIU)

subcapitol(doc, "a", "Definirea problemei",
    "Analiza statistică completă a indicatorilor de performanță: statistici "
    "descriptive, distribuții de frecvență, matrice de corelații și modele "
    "de regresie liniară multiplă OLS pentru cuantificarea factorilor determinanți.")
subcapitol(doc, "b", "Informații necesare",
    "work.financiar_macro: panel n=5 ani × 58 variabile (date financiare + macro). "
    "work.centrale_enriched: secțiune transversală n=30 centrale × 20 variabile.")
subcapitol(doc, "c", "Metode de calcul",
    "PROC MEANS: n, mean, median, std, min, max, cv, skewness, kurtosis. "
    "PROC FREQ: tabele de frecvență univariate + tabele de contingență cu chi²-test. "
    "PROC CORR: coeficienți Pearson și Spearman, matrice grafică (PLOTS=MATRIX). "
    "PROC REG: OLS cu VIF (detecție multicolinearitate), TOL, R (reziduale), "
    "CLB (intervale de încredere 95%), PLOTS=(fitplot residuals).")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "PROC MEANS — venituri_totale: medie=9.37 mld, std=2.07 mld, CV=22.1%, "
    "asimetrie=0.78 (distribuție ușor dreaptă). "
    "PROC FREQ — tip centrală: Acumulare=12 (40%), Firul apei=16 (53%), Fluvial=2 (7%). "
    "PROC CORR — r(profit_net, productie_hidro)=0.96 (Pearson), 1.00 (Spearman). "
    "PROC REG (n=30): productie_gwh_an = -4.820 + 2.267×putere_mw; R²=0.961, p<0.0001.")
subcapitol(doc, "e", "Interpretarea economică",
    "Corelația perfectă Spearman (rho=1.00) între producția hidro și profit confirmă "
    "că rangul anilor din punct de vedere al profitului este identic cu rangul "
    "din punct de vedere al producției — hidraulicitatea determină în totalitate "
    "ordinea de performanță. VIF pentru predictorii regresiei multiple pe n=5 "
    "nu este calculabil robust (n < k+1 pentru modele cu 4+ predictori), "
    "de aceea regresia multiplă se aplică pe n=30 centrale.")

doc.add_paragraph()

# ── F10 SAS ──────────────────────────────────────────────────
functie_header(doc, "SAS-10", "Grafice — PROC SGPLOT / PROC SGPANEL", PORTOCALIU)

subcapitol(doc, "a", "Definirea problemei",
    "Vizualizarea grafică a tendințelor financiare, corelației hidrologie-profit, "
    "distribuției puterii instalate și structurii cheltuielilor pentru comunicarea "
    "eficientă a rezultatelor analizei.")
subcapitol(doc, "b", "Informații necesare",
    "work.financiar_macro, work.centrale_enriched, work.financiar_indexat, "
    "work.cheltuieli_long (format long creat din ARRAY pentru SGPANEL).")
subcapitol(doc, "c", "Metode de calcul",
    "PROC SGPLOT cu: VBAR (bare verticale grupate), SERIES (linii cu markeri), "
    "SCATTER (nori de puncte), REG (linie OLS automată), VBOX (boxplot), "
    "HIGHLOW (lollipop charts). Y2AXIS pentru axă secundară. "
    "PROC SGPANEL cu PANELBY și COLUMNS/ROWS pentru grafice multiple (small multiples). "
    "ODS GRAPHICS ON / OFF, WIDTH, HEIGHT pentru controlul dimensiunilor.")
subcapitol(doc, "d", "Prezentarea rezultatelor",
    "5 grafice generate: (1) Bar+Linie: evoluție venituri, profit, marjă EBITDA 2021–2025. "
    "(2) Scatter+REG: corelație precipitații → profit net (linie trend roșie). "
    "(3) Boxplot+Jitter: distribuție putere pe tip centrală (medianele: acumulare=216, "
    "firul apei=25, fluvial=660 MW). "
    "(4) Lollipop: Top 15 centrale după producție (Porțile de Fier I: 5.200 GWh). "
    "(5) SGPANEL 3×2: evoluția celor 5 categorii de cheltuieli ca % din venituri.")
subcapitol(doc, "e", "Interpretarea economică",
    "Graficul scatter precipitații–profit relevă vizual corelația puternică (R²≈0.92 pe "
    "5 puncte): fiecare punct al indicelui de precipitații (față de baza 100) "
    "corespunde cu ~55–60 mil. RON profit suplimentar. SGPANEL cheltuieli evidențiază "
    "trendul alarmant: energia achiziționată (roșu) crește de la 1.4% la 9.5% din venituri "
    "— cel mai volatil cost și principalul risc structural al modelului de afaceri.")

add_page_break(doc)


# ═══════════════════════════════════════════════════════════════
# CONCLUZII
# ═══════════════════════════════════════════════════════════════
heading1(doc, "CONCLUZII", ALBASTRU_D)

body(doc,
     "Analiza strategică a activității Hidroelectrica S.A. în perioada 2021–2025, "
     "realizată cu Python (Streamlit) și SAS, relevă o companie cu performanțe "
     "financiare excepționale dar cu vulnerabilități structurale identificabile.")

heading2(doc, "Concluzii privind performanța financiară")
body(doc,
     "Vârful de performanță a fost înregistrat în 2023: venituri de 12.16 mld RON "
     "(+87% față de 2021), marjă EBITDA de 65.8% și profit net de 6.37 mld RON. "
     "Această performanță a fost determinată de confluența a trei factori favorabili: "
     "(1) hidraulicitate record (index 118), (2) prețuri ridicate ale energiei, "
     "(3) expansiunea portofoliului de clienți retail.", indent=True)
body(doc,
     "2025 marchează o deteriorare semnificativă: marjă EBITDA de 47.8% (cea mai "
     "slabă din perioadă), cauzată de secetă hidrologică (index 71) și creșterea "
     "energiei achiziționate de la 90 mil. RON (2021) la 911 mil. RON.", indent=True)

heading2(doc, "Concluzii privind potențialul de extindere")
body(doc,
     "Analiza clusterelor K-Means și a regresiei OLS pe n=30 centrale confirmă "
     "că puterea instalată este principalul determinant al producției (β=2.27 GWh/MW, "
     "R²=0.961). Extinderea cu 500 MW solar și 200 MW eolian ar genera estimativ "
     "+744 GWh/an producție și +357 mil. RON venituri adiționale — reducând "
     "dependența de hidraulicitate cu ~5 puncte procentuale.", indent=True)
body(doc,
     "Parteneriatul cu MASDAR (Abu Dhabi) și parcul eolian Crucea Nord "
     "(operațional 2024) confirmă direcția strategică corectă. "
     "Hidroelectrica dispune de capacitatea financiară necesară: "
     "cash din exploatare de 4.3–6.4 mld RON/an.", indent=True)

heading2(doc, "Limitele analizei")
body(doc,
     "n=5 ani pentru analiza panel (date financiare) limitează robustețea "
     "statistică a regresiei multiple și corelațiilor; n=30 centrale pentru "
     "analiza cross-secțională oferă rezultate mai fiabile. "
     "Prognozele 2026–2027 se bazează pe tendința liniară și nu incorporează "
     "scenarii hidrologice sau fluctuații de preț.", indent=True)

doc.add_paragraph()

# ── Tabel sintetic final ──────────────────────────────────────
heading2(doc, "Sinteza indicatorilor cheie Hidroelectrica 2021–2025")
tabel_simplu(doc,
    ["An", "Venituri (mld RON)", "Profit Net (mld RON)", "Marjă EBITDA (%)", "ROE (%)", "Ind. Precip."],
    [
        ["2021", "6.49", "3.12", "69.6%", "16.2%", "98 (Normal)"],
        ["2022", "9.45", "4.46", "63.3%", "20.7%", "72 (Secetos)"],
        ["2023", "12.16", "6.37", "65.8%", "25.4%", "118 (Ploios)"],
        ["2024", "9.12", "4.13", "60.2%", "18.0%", "88 (Sub-normal)"],
        ["2025", "9.62", "3.37", "47.8%", "15.2%", "71 (Secetos)"],
    ],
    col_widths=[2.5, 3.5, 3.5, 3.5, 2.5, 3.5]
)

doc.add_paragraph()

# ── Footer informativ ──────────────────────────────────────────
p_footer = doc.add_paragraph()
p_footer.paragraph_format.space_before = Pt(20)
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = p_footer.add_run(
    "Surse de date: Rapoarte anuale Hidroelectrica S.A. · BVB (H2O.RO) · "
    "ANRE · sistemulenergetic.ro · Open-Meteo Archive API · ANM\n"
    f"Document generat automat · {datetime.date.today().strftime('%d %B %Y')}"
)
rf.font.size = Pt(9)
rf.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
rf.italic = True


# ═══════════════════════════════════════════════════════════════
# SALVARE
# ═══════════════════════════════════════════════════════════════
output_path = "Proiect_Pachete_Software_Hidroelectrica_Dumitrescu.docx"
doc.save(output_path)
print(f"OK Document salvat: {output_path}")
print(f"   Sectiuni: Introducere + 9 functii Python + 10 functii SAS + Concluzii")
